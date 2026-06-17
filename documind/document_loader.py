import tempfile
from pathlib import Path
from typing import List, Optional


class Document:
    def __init__(self, content: str, metadata: dict):
        self.content = content
        self.metadata = metadata


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}
    OCR_THRESHOLD_CHARS = 50

    @classmethod
    def load(cls, file_path: str, force_ocr: bool = False) -> List[Document]:
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {cls.SUPPORTED_EXTENSIONS}")

        if ext == ".pdf":
            result = cls._load_pdf(file_path)
            if force_ocr or cls._is_scanned(result):
                from documind.ocr import OCREngine
                ocr = OCREngine()
                result = ocr.ocr_pdf(file_path)
            return result

        loader_map = {
            ".txt": cls._load_text,
            ".md": cls._load_text,
            ".docx": cls._load_docx,
        }
        return loader_map[ext](file_path)

    @classmethod
    def load_bytes(cls, filename: str, data: bytes, force_ocr: bool = False) -> List[Document]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            return cls.load(tmp_path, force_ocr=force_ocr)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    @staticmethod
    def _is_scanned(documents: List[Document]) -> bool:
        total_chars = sum(len(d.content) for d in documents)
        return total_chars < DocumentLoader.OCR_THRESHOLD_CHARS

    @staticmethod
    def _load_pdf(file_path: str) -> List[Document]:
        import pdfplumber
        docs = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        content=text.strip(),
                        metadata={"source": file_path, "page": i + 1, "type": "pdf"},
                    ))
        if not docs:
            docs.append(Document(
                content="",
                metadata={"source": file_path, "page": 0, "type": "pdf"},
            ))
        return docs

    @staticmethod
    def _load_text(file_path: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return [Document(
            content=content.strip(),
            metadata={"source": file_path, "type": "text"},
        )]

    @staticmethod
    def _load_docx(file_path: str) -> List[Document]:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        return [Document(
            content=content,
            metadata={"source": file_path, "type": "docx"},
        )]
